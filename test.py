from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import pandas as pd
import tempfile
import os

def generate_codebook(df, column_types, variable_names, category_definitions, code_df=None, output_path="codebook.docx", preview_mode=False):
    if output_path is None:
        output_path = "codebook.docx"

    doc = Document()
    doc.add_heading("Codebook Summary Report", level=1)

    # 🔹 準備欄位列表
    code_df = code_df[~code_df["Type"].astype(str).str.lower().eq("0")]
    
    # 🔹 遺失值統計區塊
    doc.add_heading("Missing Value Summary", level=2)
    na_counts = df.isnull().sum()
    na_percent = df.isnull().mean() * 100

    na_df = pd.DataFrame({
        "Column": na_counts.index,
        "Missing Count": na_counts.values,
        "Missing Rate (%)": na_percent.round(2).values
    }).query("`Missing Count` > 0").reset_index(drop=True)

    if not na_df.empty:
        table = doc.add_table(rows=1 + len(na_df), cols=3)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Column"
        table.cell(0, 1).text = "Missing Count"
        table.cell(0, 2).text = "Missing Rate (%)"
        for i, row in na_df.iterrows():
            table.cell(i + 1, 0).text = str(row["Column"])
            table.cell(i + 1, 1).text = str(row["Missing Count"])
            table.cell(i + 1, 2).text = str(row["Missing Rate (%)"])
    else:
        doc.add_paragraph("No missing values in any columns.")

    # 🔹 變數類型統計區塊
    doc.add_heading("Variable Type Summary", level=2)
    type_count = pd.Series(column_types).value_counts().sort_index()
    type_label_map = {1: "數值型 (Type 1)", 2: "類別型 (Type 2)"}

    table = doc.add_table(rows=1 + len(type_count), cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "變數類型"
    table.cell(0, 1).text = "欄位數"
    for i, (type_code, count) in enumerate(type_count.items()):
        label = type_label_map.get(type_code, f"其他 ({type_code})")
        table.cell(i + 1, 0).text = label
        table.cell(i + 1, 1).text = str(count)

    # 🔹 開始處理欄位細節
    df = df.dropna()  # ✅ 去除有缺失的 row
    columns = code_df["Column"] if code_df is not None else df.columns

    for col in columns:
        if col not in column_types:
            continue
        type_code = column_types[col]
        if type_code == 0:
            continue

        var_name = variable_names.get(col, col)
        doc.add_heading(f"Variable: {col} ({var_name})", level=2)

        # 🟦 類別型變數
        if type_code == 2:
            value_counts = df[col].value_counts(dropna=False).sort_index()
            total = len(df)
            defs = category_definitions.get(col, {})
            lines = [
                f"{k}: {defs.get(k, '')} → {v} ({v/total:.2%})"
                for k, v in value_counts.items()
            ]
            summary_text = "\n".join(lines)

            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Variable Name"
            table.cell(0, 1).text = f"{col} ({var_name})"
            table.cell(1, 0).text = "Categories Summary"
            table.cell(1, 1).text = summary_text

            # 圖表
            fig, ax = plt.subplots()
            value_counts.plot(kind="bar", color="cornflowerblue", ax=ax)
            ax.set_title(f"Count Plot of {col}")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            plt.tight_layout()
            plt.savefig(tmp.name)
            plt.close()
            doc.add_picture(tmp.name, width=Inches(4.5))
            try: os.unlink(tmp.name)
            except PermissionError: pass

        # 🟩 數值型變數
        elif type_code == 1:
            try:
                df[col] = pd.to_numeric(df[col], errors="coerce")
            except Exception:
                continue
            if df[col].dropna().empty:
                continue

            desc = df[col].describe()
            table = doc.add_table(rows=4, cols=4)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Index"
            table.cell(0, 1).text = col
            table.cell(0, 2).text = "Variable Name"
            table.cell(0, 3).text = var_name

            table.cell(1, 0).text = "Mean"
            table.cell(1, 1).text = f"{desc['mean']:.3f}"
            table.cell(1, 2).text = "Std Dev"
            table.cell(1, 3).text = f"{desc['std']:.3f}"

            table.cell(2, 0).text = "Max"
            table.cell(2, 1).text = f"{desc['max']:.3f}"
            table.cell(2, 2).text = "Min"
            table.cell(2, 3).text = f"{desc['min']:.3f}"

            table.cell(3, 0).text = "Q1 (25%)"
            table.cell(3, 1).text = f"{desc['25%']:.3f}"
            table.cell(3, 2).text = "Q3 (75%)"
            table.cell(3, 3).text = f"{desc['75%']:.3f}"

            # Histogram
            fig, ax = plt.subplots()
            df[col].plot(kind="hist", bins=10, color="skyblue", edgecolor="black", ax=ax)
            ax.set_title(f"Histogram of {col}")
            tmp1 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            plt.tight_layout()
            plt.savefig(tmp1.name)
            plt.close()
            doc.add_picture(tmp1.name, width=Inches(4.5))
            try: os.unlink(tmp1.name)
            except PermissionError: pass

            # Boxplot
            fig2, ax2 = plt.subplots()
            df.boxplot(column=col, ax=ax2)
            ax2.set_title(f"Boxplot of {col}")
            tmp2 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            plt.tight_layout()
            plt.savefig(tmp2.name)
            plt.close()
            doc.add_picture(tmp2.name, width=Inches(4.5))
            try: os.unlink(tmp2.name)
            except PermissionError: pass

    doc.save(output_path)
    return output_path
